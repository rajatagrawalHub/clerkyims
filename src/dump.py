import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Function to add code with black background and white text
def add_code_block(paragraph, code_text):
    run = paragraph.add_run(code_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)

    # Apply black background and white text
    highlight = run._element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '000000')  # Black background
    highlight.rPr.append(shd)
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text

# Create a new Word document
doc = Document()

# Set your target directory here
base_path = 'src'  # <-- Replace this with the actual path

# Walk through the directory and its subdirectories
for root, dirs, files in os.walk(base_path):
    for file in files:
        if file.endswith(('.js', '.jsx', '.css')):
            file_path = os.path.join(root, file)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # Add file name as title
                title = doc.add_paragraph()
                run = title.add_run(file)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13.5)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # One-line gap
                doc.add_paragraph()

                # Add code content
                code_paragraph = doc.add_paragraph()
                add_code_block(code_paragraph, content)

                # Add extra space between files
                doc.add_paragraph()
                doc.add_paragraph()

            except Exception as e:
                print(f"Error reading {file_path}: {e}")

# Save the document
doc.save('code_dump.docx')
print("Document created: jsx_js_code_dump.docx")
